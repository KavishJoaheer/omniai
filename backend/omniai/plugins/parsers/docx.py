from __future__ import annotations

import io

from docx import Document as DocxDocument

from omniai.ports.parser import ParseResult


class DocxParser:
    name = "docx"
    mime_types: tuple[str, ...] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    extensions: tuple[str, ...] = (".docx",)

    def parse(self, *, data: bytes, filename: str) -> ParseResult:
        doc = DocxDocument(io.BytesIO(data))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return ParseResult(
            text="\n".join(parts),
            page_count=1,
            metadata={"filename": filename, "paragraph_count": len(doc.paragraphs)},
        )
